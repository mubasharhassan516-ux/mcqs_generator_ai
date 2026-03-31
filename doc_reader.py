"""
doc_reader.py
=============
Extracts text from Word (.docx), PDF (.pdf), and text files (.txt, .rtf, .md).
Detects chapters/sections and chunks content proportionally.
"""

import re
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
#  TEXT FILE READER (New)
# ─────────────────────────────────────────────

def extract_from_text(file_path: str) -> list[dict]:
    """
    Extract text from a .txt, .rtf, or .md file.
    Attempts to detect chapters based on common patterns.
    """
    try:
        # Try UTF-8 first
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        # Fall back to latin-1
        with open(file_path, 'r', encoding='latin-1') as f:
            content = f.read()
    
    chapters = []
    lines = content.split('\n')
    current_chapter = "Introduction"
    current_text = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if line looks like a chapter heading
        if _looks_like_chapter_heading(line) or _looks_like_markdown_heading(line):
            if current_text:
                chapters.append({
                    "chapter": current_chapter,
                    "text": " ".join(current_text).strip()
                })
            current_chapter = line
            current_text = []
        else:
            current_text.append(line)
    
    # Add last chapter
    if current_text:
        chapters.append({
            "chapter": current_chapter,
            "text": " ".join(current_text).strip()
        })
    
    if not chapters:
        # If no chapters detected, treat entire file as one chapter
        chapters = [{
            "chapter": Path(file_path).stem,
            "text": content
        }]
    
    return chapters


def _looks_like_markdown_heading(line: str) -> bool:
    """Detect Markdown-style headings (# Heading, ## Heading, etc.)"""
    return line.startswith('#') and len(line) > 1 and line[1] != '#'


# ─────────────────────────────────────────────
#  WORD DOCUMENT READER
# ─────────────────────────────────────────────

def extract_from_docx(file_path: str) -> list[dict]:
    """
    Extract text from a .docx file and detect chapters via heading styles.
    Returns a list of dicts: [{"chapter": str, "text": str}, ...]
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    doc = Document(file_path)
    chapters = []
    current_chapter = "Introduction"
    current_text = []

    heading_styles = {"Heading 1", "Heading 2", "Heading 3",
                      "Title", "Subtitle"}

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            continue

        # Detect a chapter heading
        if style_name in heading_styles or _looks_like_chapter_heading(text):
            # Save the previous chapter
            if current_text:
                chapters.append({
                    "chapter": current_chapter,
                    "text": " ".join(current_text).strip()
                })
            current_chapter = text
            current_text = []
        else:
            current_text.append(text)

    # Save the last chapter
    if current_text:
        chapters.append({
            "chapter": current_chapter,
            "text": " ".join(current_text).strip()
        })

    if not chapters:
        raise ValueError("No readable text found in the Word document.")

    return chapters


# ─────────────────────────────────────────────
#  PDF READER
# ─────────────────────────────────────────────

def extract_from_pdf(file_path: str) -> list[dict]:
    """
    Extract text from a .pdf file using pdfplumber.
    Detects chapters by scanning lines that look like headings.
    Returns a list of dicts: [{"chapter": str, "text": str}, ...]
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required. Install with: pip install pdfplumber")

    chapters = []
    current_chapter = "Introduction"
    current_text = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if not page_text:
                continue

            for line in page_text.splitlines():
                line = line.strip()
                if not line:
                    continue

                if _looks_like_chapter_heading(line):
                    if current_text:
                        chapters.append({
                            "chapter": current_chapter,
                            "text": " ".join(current_text).strip()
                        })
                    current_chapter = line
                    current_text = []
                else:
                    current_text.append(line)

    # Save the final chapter
    if current_text:
        chapters.append({
            "chapter": current_chapter,
            "text": " ".join(current_text).strip()
        })

    if not chapters:
        raise ValueError("No readable text found in the PDF.")

    return chapters


# ─────────────────────────────────────────────
#  CHAPTER HEADING DETECTOR
# ─────────────────────────────────────────────

def _looks_like_chapter_heading(text: str) -> bool:
    """
    Heuristic: a line is a chapter heading if it:
    - Starts with 'Chapter', 'Section', 'Unit', 'Part', 'Module'
    - Matches patterns like '1.', '1.1', 'I.', 'A.'
    - Is short (≤ 80 chars) and title-cased or ALL CAPS
    """
    if len(text) > 120:
        return False

    chapter_keywords = re.compile(
        r'^(chapter|section|unit|part|module|lesson|topic|appendix)\b',
        re.IGNORECASE
    )
    numbered = re.compile(r'^(\d+[\.\)]\s|\d+\.\d+[\.\)]\s|[IVXLC]+[\.\)]\s|[A-Z][\.\)]\s)')

    if chapter_keywords.match(text):
        return True
    if numbered.match(text):
        return True
    # Short, ALL CAPS line
    if text.isupper() and 5 < len(text) < 80:
        return True

    return False


# ─────────────────────────────────────────────
#  DIRECT TEXT EXTRACTION FUNCTIONS (for app.py)
# ─────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """Extract raw text from PDF file (for direct use, not chapterized)"""
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text from DOCX file (for direct use, not chapterized)"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """Extract raw text from TXT file (for direct use)"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        return ""


# ─────────────────────────────────────────────
#  MAIN ENTRY POINT
# ─────────────────────────────────────────────

def extract_chapters(file_path: str) -> list[dict]:
    """
    Main dispatcher: auto-detects file type and extracts chapters.
    Returns: [{"chapter": str, "text": str}, ...]
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        chapters = extract_from_docx(file_path)
    elif suffix == ".pdf":
        chapters = extract_from_pdf(file_path)
    elif suffix in [".txt", ".rtf", ".md"]:
        chapters = extract_from_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: .docx, .pdf, .txt, .rtf, .md")

    # Filter out chapters with too little text
    chapters = [c for c in chapters if len(c["text"].split()) >= 20]

    logger.info(f"Extracted {len(chapters)} chapter(s) from {path.name}")
    return chapters


def distribute_mcq_count(chapters: list[dict], total_mcqs: int) -> list[dict]:
    """
    Distribute MCQ count across chapters proportionally by word count.
    Returns chapters with an added 'mcq_count' key.
    """
    total_words = sum(len(c["text"].split()) for c in chapters)
    if total_words == 0:
        raise ValueError("Document appears to have no readable content.")

    remaining = total_mcqs
    for i, chapter in enumerate(chapters):
        word_count = len(chapter["text"].split())
        if i == len(chapters) - 1:
            # Give remaining MCQs to the last chapter to avoid rounding loss
            chapter["mcq_count"] = max(1, remaining)
        else:
            count = max(1, round(total_mcqs * word_count / total_words))
            count = min(count, remaining - (len(chapters) - i - 1))
            chapter["mcq_count"] = max(1, count)
            remaining -= chapter["mcq_count"]

    return chapters